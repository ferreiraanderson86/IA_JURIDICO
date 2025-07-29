<<<<<<< HEAD
import os
import argparse
import pdfplumber
import docx
from dotenv import load_dotenv
from openai import OpenAI

# ---------------------------
# ConfiguraÃ§Ãµes
# ---------------------------

SYSTEM_PROMPT_ANALISE = (
    "VocÃª Ã© um advogado brasileiro especialista em direito do consumidor e bancÃ¡rio. "
    "Analise o contrato enviado, identifique clÃ¡usulas potencialmente abusivas, riscos, "
    "fundamentaÃ§Ã£o jurÃ­dica (CDC, STJ, SÃºmulas, artigos do CC, etc.), calcule/indique possÃ­veis "
    "vantagens econÃ´micas e sugira estratÃ©gias (administrativas e judiciais). "
    "Seja claro, objetivo, estruturado e aponte artigos/sÃºmulas quando pertinente."
)

USER_PROMPT_TEMPLATE = """
Analise o seguinte contrato bancÃ¡rio (ou de prestaÃ§Ã£o de serviÃ§os) e elabore um parecer jurÃ­dico completo contendo:

1) **Resumo do contrato** (partes, objeto, duraÃ§Ã£o, valores, multas, foro, etc.)
2) **Pontos crÃ­ticos / clÃ¡usulas potencialmente abusivas** (explique o porquÃª e cite fundamento legal/jurisprudencial)
3) **Riscos ao contratante**
4) **FundamentaÃ§Ã£o jurÃ­dica aplicÃ¡vel** (CDC, CC, STJ, sÃºmulas, resoluÃ§Ãµes do BACEN, etc.)
5) **EstratÃ©gias e pedidos possÃ­veis** (administrativas e judiciais)
6) **Documentos e provas que devem ser reunidos**
7) **Estimativa de vantagem econÃ´mica (se aplicÃ¡vel)**
8) **Modelo de pedidos (tÃ³picos)**

Contrato (trecho ou inteiro):
\"\"\"{texto}\"\"\"
"""

SYSTEM_PROMPT_PETICAO = (
    "VocÃª Ã© um advogado brasileiro especialista em petiÃ§Ãµes iniciais. "
    "Com base na anÃ¡lise jurÃ­dica fornecida, redija uma petiÃ§Ã£o inicial completa, "
    "bem estruturada, com linguagem tÃ©cnica adequada, incluindo endereÃ§amento, qualificaÃ§Ã£o das partes, "
    "fatos, fundamentos jurÃ­dicos (com artigos, sÃºmulas, precedentes), pedidos (com tutela de urgÃªncia se couber), "
    "valor da causa, provas, e requerimentos finais."
)

USER_PROMPT_PETICAO = """
Use a anÃ¡lise jurÃ­dica abaixo para redigir a **petiÃ§Ã£o inicial completa**.
Inclua: endereÃ§amento, qualificaÃ§Ã£o das partes (coloque campos para serem preenchidos), fatos, fundamentos jurÃ­dicos
(detalhados com artigos, sÃºmulas e precedentes), pedidos, valor da causa (campo para preenchimento), provas, e requerimentos finais.
Se pertinente, inclua pedidos de tutela de urgÃªncia.

AnÃ¡lise jurÃ­dica:
\"\"\"{analise}\"\"\"
"""

# ---------------------------
# Leitura de arquivos
# ---------------------------

def ler_pdf(caminho: str) -> str:
    with pdfplumber.open(caminho) as pdf:
        textos = []
        for page in pdf.pages:
            txt = page.extract_text() or ""
            textos.append(txt)
    return "\n".join(textos)

def ler_docx(caminho: str) -> str:
    d = docx.Document(caminho)
    return "\n".join(p.text for p in d.paragraphs)

def ler_arquivo(caminho: str) -> str:
    ext = os.path.splitext(caminho.lower())[1]
    if ext == ".pdf":
        return ler_pdf(caminho)
    elif ext == ".docx":
        return ler_docx(caminho)
    else:
        raise ValueError("Formato nÃ£o suportado. Use PDF ou DOCX.")

# ---------------------------
# OpenAI helpers
# ---------------------------

def get_client():
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY nÃ£o encontrada. Defina no .env.")
    return OpenAI(api_key=api_key)

def call_chat_completion(client: OpenAI, system_prompt: str, user_prompt: str, model: str = "gpt-4o"):
    return client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    ).choices[0].message.content

# ---------------------------
# LÃ³gica principal
# ---------------------------

def analisar_contrato(caminho_arquivo: str, model: str = "gpt-4o") -> str:
    texto = ler_arquivo(caminho_arquivo)

    # Se o contrato for muito grande, vocÃª pode optar por truncar ou fatiar.
    # Aqui, deixei simples. Ajuste se necessÃ¡rio.
    prompt_user = USER_PROMPT_TEMPLATE.format(texto=texto[:120000])  # ~120k chars de seguranÃ§a

    client = get_client()
    analise = call_chat_completion(client, SYSTEM_PROMPT_ANALISE, prompt_user, model=model)
    return analise

def gerar_peticao(analise: str, model: str = "gpt-4o") -> str:
    client = get_client()
    user_prompt = USER_PROMPT_PETICAO.format(analise=analise)
    peticao = call_chat_completion(client, SYSTEM_PROMPT_PETICAO, user_prompt, model=model)
    return peticao

def salvar(texto: str, caminho: str):
    with open(caminho, "w", encoding="utf-8") as f:
        f.write(texto)

def main():
    parser = argparse.ArgumentParser(description="Analisar contrato bancÃ¡rio com GPT-4o e (opcional) gerar petiÃ§Ã£o.")
    parser.add_argument("--arquivo", "-a", required=True, help="Caminho do contrato (PDF ou DOCX)")
    parser.add_argument("--saida", "-s", default="parecer.md", help="Arquivo de saÃ­da do parecer (.md)")
    parser.add_argument("--modelo", "-m", default="gpt-4o", help="Modelo (ex: gpt-4o, gpt-4.1, gpt-4o-mini)")
    parser.add_argument("--peticao", action="store_true", help="Gerar tambÃ©m a petiÃ§Ã£o inicial com base na anÃ¡lise")
    parser.add_argument("--saida-peticao", default="peticao.md", help="Arquivo de saÃ­da da petiÃ§Ã£o (.md)")
    args = parser.parse_args()

    print("ðŸ‘‰ Lendo contrato...")
    try:
        analise = analisar_contrato(args.arquivo, model=args.modelo)
    except Exception as e:
        print("Erro ao analisar contrato:", e)
        return

    salvar(analise, args.saida)
    print(f"âœ… Parecer salvo em: {args.saida}")

    if args.peticao:
        print("ðŸ‘‰ Gerando petiÃ§Ã£o inicial...")
        try:
            peticao = gerar_peticao(analise, model=args.modelo)
            salvar(peticao, args.saida_peticao)
            print(f"âœ… PetiÃ§Ã£o salva em: {args.saida_peticao}")
        except Exception as e:
            print("Erro ao gerar petiÃ§Ã£o:", e)

if __name__ == "__main__":
    main()
=======
import os
import argparse
import pdfplumber
import docx
from dotenv import load_dotenv
from openai import OpenAI

# ---------------------------
# ConfiguraÃ§Ãµes
# ---------------------------

SYSTEM_PROMPT_ANALISE = (
    "VocÃª Ã© um advogado brasileiro especialista em direito do consumidor e bancÃ¡rio. "
    "Analise o contrato enviado, identifique clÃ¡usulas potencialmente abusivas, riscos, "
    "fundamentaÃ§Ã£o jurÃ­dica (CDC, STJ, SÃºmulas, artigos do CC, etc.), calcule/indique possÃ­veis "
    "vantagens econÃ´micas e sugira estratÃ©gias (administrativas e judiciais). "
    "Seja claro, objetivo, estruturado e aponte artigos/sÃºmulas quando pertinente."
)

USER_PROMPT_TEMPLATE = """
Analise o seguinte contrato bancÃ¡rio (ou de prestaÃ§Ã£o de serviÃ§os) e elabore um parecer jurÃ­dico completo contendo:

1) **Resumo do contrato** (partes, objeto, duraÃ§Ã£o, valores, multas, foro, etc.)
2) **Pontos crÃ­ticos / clÃ¡usulas potencialmente abusivas** (explique o porquÃª e cite fundamento legal/jurisprudencial)
3) **Riscos ao contratante**
4) **FundamentaÃ§Ã£o jurÃ­dica aplicÃ¡vel** (CDC, CC, STJ, sÃºmulas, resoluÃ§Ãµes do BACEN, etc.)
5) **EstratÃ©gias e pedidos possÃ­veis** (administrativas e judiciais)
6) **Documentos e provas que devem ser reunidos**
7) **Estimativa de vantagem econÃ´mica (se aplicÃ¡vel)**
8) **Modelo de pedidos (tÃ³picos)**

Contrato (trecho ou inteiro):
\"\"\"{texto}\"\"\"
"""

SYSTEM_PROMPT_PETICAO = (
    "VocÃª Ã© um advogado brasileiro especialista em petiÃ§Ãµes iniciais. "
    "Com base na anÃ¡lise jurÃ­dica fornecida, redija uma petiÃ§Ã£o inicial completa, "
    "bem estruturada, com linguagem tÃ©cnica adequada, incluindo endereÃ§amento, qualificaÃ§Ã£o das partes, "
    "fatos, fundamentos jurÃ­dicos (com artigos, sÃºmulas, precedentes), pedidos (com tutela de urgÃªncia se couber), "
    "valor da causa, provas, e requerimentos finais."
)

USER_PROMPT_PETICAO = """
Use a anÃ¡lise jurÃ­dica abaixo para redigir a **petiÃ§Ã£o inicial completa**.
Inclua: endereÃ§amento, qualificaÃ§Ã£o das partes (coloque campos para serem preenchidos), fatos, fundamentos jurÃ­dicos
(detalhados com artigos, sÃºmulas e precedentes), pedidos, valor da causa (campo para preenchimento), provas, e requerimentos finais.
Se pertinente, inclua pedidos de tutela de urgÃªncia.

AnÃ¡lise jurÃ­dica:
\"\"\"{analise}\"\"\"
"""

# ---------------------------
# Leitura de arquivos
# ---------------------------

def ler_pdf(caminho: str) -> str:
    with pdfplumber.open(caminho) as pdf:
        textos = []
        for page in pdf.pages:
            txt = page.extract_text() or ""
            textos.append(txt)
    return "\n".join(textos)

def ler_docx(caminho: str) -> str:
    d = docx.Document(caminho)
    return "\n".join(p.text for p in d.paragraphs)

def ler_arquivo(caminho: str) -> str:
    ext = os.path.splitext(caminho.lower())[1]
    if ext == ".pdf":
        return ler_pdf(caminho)
    elif ext == ".docx":
        return ler_docx(caminho)
    else:
        raise ValueError("Formato nÃ£o suportado. Use PDF ou DOCX.")

# ---------------------------
# OpenAI helpers
# ---------------------------

def get_client():
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY nÃ£o encontrada. Defina no .env.")
    return OpenAI(api_key=api_key)

def call_chat_completion(client: OpenAI, system_prompt: str, user_prompt: str, model: str = "gpt-4o"):
    return client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    ).choices[0].message.content

# ---------------------------
# LÃ³gica principal
# ---------------------------

def analisar_contrato(caminho_arquivo: str, model: str = "gpt-4o") -> str:
    texto = ler_arquivo(caminho_arquivo)

    # Se o contrato for muito grande, vocÃª pode optar por truncar ou fatiar.
    # Aqui, deixei simples. Ajuste se necessÃ¡rio.
    prompt_user = USER_PROMPT_TEMPLATE.format(texto=texto[:120000])  # ~120k chars de seguranÃ§a

    client = get_client()
    analise = call_chat_completion(client, SYSTEM_PROMPT_ANALISE, prompt_user, model=model)
    return analise

def gerar_peticao(analise: str, model: str = "gpt-4o") -> str:
    client = get_client()
    user_prompt = USER_PROMPT_PETICAO.format(analise=analise)
    peticao = call_chat_completion(client, SYSTEM_PROMPT_PETICAO, user_prompt, model=model)
    return peticao

def salvar(texto: str, caminho: str):
    with open(caminho, "w", encoding="utf-8") as f:
        f.write(texto)

def main():
    parser = argparse.ArgumentParser(description="Analisar contrato bancÃ¡rio com GPT-4o e (opcional) gerar petiÃ§Ã£o.")
    parser.add_argument("--arquivo", "-a", required=True, help="Caminho do contrato (PDF ou DOCX)")
    parser.add_argument("--saida", "-s", default="parecer.md", help="Arquivo de saÃ­da do parecer (.md)")
    parser.add_argument("--modelo", "-m", default="gpt-4o", help="Modelo (ex: gpt-4o, gpt-4.1, gpt-4o-mini)")
    parser.add_argument("--peticao", action="store_true", help="Gerar tambÃ©m a petiÃ§Ã£o inicial com base na anÃ¡lise")
    parser.add_argument("--saida-peticao", default="peticao.md", help="Arquivo de saÃ­da da petiÃ§Ã£o (.md)")
    args = parser.parse_args()

    print("ðŸ‘‰ Lendo contrato...")
    try:
        analise = analisar_contrato(args.arquivo, model=args.modelo)
    except Exception as e:
        print("Erro ao analisar contrato:", e)
        return

    salvar(analise, args.saida)
    print(f"âœ… Parecer salvo em: {args.saida}")

    if args.peticao:
        print("ðŸ‘‰ Gerando petiÃ§Ã£o inicial...")
        try:
            peticao = gerar_peticao(analise, model=args.modelo)
            salvar(peticao, args.saida_peticao)
            print(f"âœ… PetiÃ§Ã£o salva em: {args.saida_peticao}")
        except Exception as e:
            print("Erro ao gerar petiÃ§Ã£o:", e)

if __name__ == "__main__":
    main()
>>>>>>> 84bc48d03e12871b7a7bf01d0c98fbf1f1bbbe91
