<<<<<<< HEAD
import streamlit as st
import tempfile
import os
from docx import Document
from analisador import ler_arquivo, gerar_analise, gerar_resposta_chat

st.set_page_config(page_title="Plataforma JurÃ­dica IA", layout="wide")
st.title("âš–ï¸ Plataforma JurÃ­dica com IA")

# ------------------------------
# Session State
# ------------------------------
if "analises" not in st.session_state:
    st.session_state.analises = {}

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "area" not in st.session_state:
    st.session_state.area = "Civil"

# ------------------------------
# Ãrea JurÃ­dica
# ------------------------------
area = st.selectbox("Escolha a Ã¡rea jurÃ­dica:", ["Civil", "Criminal", "BancÃ¡rio", "TributÃ¡rio"])
st.session_state.area = area

def get_prompt_sistema(area: str) -> str:
    base = (
        "VocÃª Ã© um advogado brasileiro experiente, redator tÃ©cnico jurÃ­dico, "
        "especializado em {area}. VocÃª deve responder de forma clara, fundamentada, precisa e formal, "
        "utilizando linguagem adequada ao meio jurÃ­dico e considerando as normas vigentes no Brasil. "
        "Sempre que possÃ­vel, fundamente suas respostas com base em legislaÃ§Ã£o, jurisprudÃªncia atualizada, "
        "doutrina majoritÃ¡ria e boas prÃ¡ticas processuais."
    )

    if area == "Civil":
        return base.format(area="Direito Civil") + (
            " Foque no CÃ³digo Civil, CÃ³digo de Processo Civil, enunciados das Jornadas de Direito Civil do CJF, "
            "e jurisprudÃªncia do STJ e STF. Considere contratos, obrigaÃ§Ãµes, responsabilidade civil, famÃ­lia e sucessÃµes."
        )

    elif area == "Criminal":
        return base.format(area="Direito Penal") + (
            " Utilize o CÃ³digo Penal, CÃ³digo de Processo Penal, jurisprudÃªncia do STJ e STF, "
            "incluindo sÃºmulas e precedentes relevantes. Fundamente defesas, recursos, habeas corpus e medidas cautelares."
        )

    elif area == "BancÃ¡rio":
        return base.format(area="Direito BancÃ¡rio e do Consumidor") + (
            " Baseie-se no CÃ³digo de Defesa do Consumidor (CDC), resoluÃ§Ãµes do Banco Central (BACEN), "
            "jurisprudÃªncia do STJ, contratos bancÃ¡rios e prÃ¡ticas abusivas. Considere revisÃ£o de clÃ¡usulas, juros abusivos "
            "e aÃ§Ãµes revisionais de contrato."
        )

    elif area == "TributÃ¡rio":
        return base.format(area="Direito TributÃ¡rio") + (
            " Fundamente suas respostas no CÃ³digo TributÃ¡rio Nacional (CTN), ConstituiÃ§Ã£o Federal, "
            "legislaÃ§Ã£o infraconstitucional, decisÃµes do CARF, e jurisprudÃªncia do STJ e STF. "
            "Considere tambÃ©m princÃ­pios como legalidade, anterioridade e capacidade contributiva."
        )

    return base.format(area="Direito") + " Atue como um advogado generalista bem preparado e objetivo."


# Inicializa histÃ³rico com prompt da Ã¡rea
if not st.session_state.chat_history:
    st.session_state.chat_history.append({"role": "system", "content": get_prompt_sistema(area)})

# Se jÃ¡ houver anÃ¡lises feitas, adicione como contexto
if st.session_state.analises:
    for nome, analise in st.session_state.analises.items():
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": (
                f"AnÃ¡lise anterior do documento **{nome}**:\n\n"
                f"{analise[:3000]}"
            )
        })

# ------------------------------
# Upload de Arquivos
# ------------------------------
uploaded_files = st.file_uploader(
    "ðŸ“Ž Envie um ou mais arquivos (PDF ou DOCX)", type=["pdf", "docx"], accept_multiple_files=True
)

col1, col2 = st.columns([1, 1])

# ------------------------------
# Analisar contratos (Civil e BancÃ¡rio)
# ------------------------------
if area in ["Civil", "BancÃ¡rio"] and uploaded_files:
    with col1:
        if st.button("ðŸ” Analisar Contratos"):
            for up in uploaded_files:
                suffix = os.path.splitext(up.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(up.read())
                    caminho_temp = tmp.name

                try:
                    with st.spinner(f"Analisando {up.name}..."):
                        texto = ler_arquivo(caminho_temp)
                        analise = gerar_analise(texto, area=area)
                        st.session_state.analises[up.name] = analise
                        st.session_state.chat_history.append(
                            {"role": "assistant", "content": f"AnÃ¡lise do contrato {up.name}:\n{analise}"}
                        )
                except Exception as e:
                    st.error(f"Erro ao analisar {up.name}: {e}")

# ------------------------------
# Limpar anÃ¡lises
# ------------------------------
with col2:
    if st.session_state.analises and st.button("ðŸ—‘ Limpar anÃ¡lises"):
        st.session_state.analises = {}

# ------------------------------
# ExibiÃ§Ã£o das anÃ¡lises
# ------------------------------
if st.session_state.analises:
    st.subheader("ðŸ“‘ AnÃ¡lises geradas")

    for nome, analise in st.session_state.analises.items():
        with st.expander(f"{nome}"):
            st.text_area("Parecer JurÃ­dico", analise, height=350)

            st.download_button(
                "â¬‡ Baixar parecer",
                analise.encode("utf-8"),
                file_name=f"parecer_{nome}.txt",
                mime="text/plain"
            )

            if st.button(f"âœ Gerar PetiÃ§Ã£o Inicial para {nome}"):
                with st.spinner("Gerando petiÃ§Ã£o inicial..."):
                    try:
                        doc = Document()
                        doc.add_heading("PetiÃ§Ã£o Inicial", 0)
                        doc.add_paragraph(analise)
                        peticao_path = os.path.join(tempfile.gettempdir(), f"peticao_inicial_{nome}.docx")
                        doc.save(peticao_path)

                        with open(peticao_path, "rb") as f:
                            st.download_button(
                                f"â¬‡ Baixar PetiÃ§Ã£o ({nome})",
                                f,
                                file_name=f"peticao_inicial_{nome}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Erro ao gerar petiÃ§Ã£o: {e}")
else:
    st.info("Nenhuma anÃ¡lise disponÃ­vel ainda.")

# ------------------------------
# Chat JurÃ­dico
# ------------------------------
st.divider()
st.subheader(f"ðŸ’¬ Chat JurÃ­dico â€“ Ãrea: {area}")

# Exibe histÃ³rico
for msg in st.session_state.chat_history:
    if msg["role"] == "user":
        with st.chat_message("user"):
            st.markdown(msg["content"])
    elif msg["role"] == "assistant":
        with st.chat_message("assistant"):
            st.markdown(msg["content"])

# Entrada do chat
if prompt := st.chat_input("Digite sua pergunta para a IA jurÃ­dica"):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("Analisando..."):
            try:
                resposta = gerar_resposta_chat(st.session_state.chat_history)
                st.markdown(resposta)
                st.session_state.chat_history.append({"role": "assistant", "content": resposta})
            except Exception as e:
                st.error(f"Erro ao responder: {e}")

# ------------------------------
# Sidebar
# ------------------------------
with st.sidebar:
    st.header("âš™ï¸ ConfiguraÃ§Ãµes")
    if st.button("ðŸ§¹ Limpar histÃ³rico do chat"):
        st.session_state.chat_history = [{"role": "system", "content": get_prompt_sistema(area)}]
    st.markdown("---")
    st.caption("Desenvolvido com â¤ï¸ por vocÃª e IA")
=======
import streamlit as st
import tempfile
import os
from docx import Document
from analisador import ler_arquivo, gerar_analise, gerar_resposta_chat

st.set_page_config(page_title="Plataforma JurÃ­dica IA", layout="wide")
st.title("âš–ï¸ Plataforma JurÃ­dica com IA")

# ------------------------------
# Session State
# ------------------------------
if "analises" not in st.session_state:
    st.session_state.analises = {}

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "area" not in st.session_state:
    st.session_state.area = "Civil"

# ------------------------------
# Ãrea JurÃ­dica
# ------------------------------
area = st.selectbox("Escolha a Ã¡rea jurÃ­dica:", ["Civil", "Criminal", "BancÃ¡rio", "TributÃ¡rio"])
st.session_state.area = area

def get_prompt_sistema(area: str) -> str:
    base = (
        "VocÃª Ã© um advogado brasileiro experiente, redator tÃ©cnico jurÃ­dico, "
        "especializado em {area}. VocÃª deve responder de forma clara, fundamentada, precisa e formal, "
        "utilizando linguagem adequada ao meio jurÃ­dico e considerando as normas vigentes no Brasil. "
        "Sempre que possÃ­vel, fundamente suas respostas com base em legislaÃ§Ã£o, jurisprudÃªncia atualizada, "
        "doutrina majoritÃ¡ria e boas prÃ¡ticas processuais."
    )

    if area == "Civil":
        return base.format(area="Direito Civil") + (
            " Foque no CÃ³digo Civil, CÃ³digo de Processo Civil, enunciados das Jornadas de Direito Civil do CJF, "
            "e jurisprudÃªncia do STJ e STF. Considere contratos, obrigaÃ§Ãµes, responsabilidade civil, famÃ­lia e sucessÃµes."
        )

    elif area == "Criminal":
        return base.format(area="Direito Penal") + (
            " Utilize o CÃ³digo Penal, CÃ³digo de Processo Penal, jurisprudÃªncia do STJ e STF, "
            "incluindo sÃºmulas e precedentes relevantes. Fundamente defesas, recursos, habeas corpus e medidas cautelares."
        )

    elif area == "BancÃ¡rio":
        return base.format(area="Direito BancÃ¡rio e do Consumidor") + (
            " Baseie-se no CÃ³digo de Defesa do Consumidor (CDC), resoluÃ§Ãµes do Banco Central (BACEN), "
            "jurisprudÃªncia do STJ, contratos bancÃ¡rios e prÃ¡ticas abusivas. Considere revisÃ£o de clÃ¡usulas, juros abusivos "
            "e aÃ§Ãµes revisionais de contrato."
        )

    elif area == "TributÃ¡rio":
        return base.format(area="Direito TributÃ¡rio") + (
            " Fundamente suas respostas no CÃ³digo TributÃ¡rio Nacional (CTN), ConstituiÃ§Ã£o Federal, "
            "legislaÃ§Ã£o infraconstitucional, decisÃµes do CARF, e jurisprudÃªncia do STJ e STF. "
            "Considere tambÃ©m princÃ­pios como legalidade, anterioridade e capacidade contributiva."
        )

    return base.format(area="Direito") + " Atue como um advogado generalista bem preparado e objetivo."


# Inicializa histÃ³rico com prompt da Ã¡rea
if not st.session_state.chat_history:
    st.session_state.chat_history.append({"role": "system", "content": get_prompt_sistema(area)})

# Se jÃ¡ houver anÃ¡lises feitas, adicione como contexto
if st.session_state.analises:
    for nome, analise in st.session_state.analises.items():
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": (
                f"AnÃ¡lise anterior do documento **{nome}**:\n\n"
                f"{analise[:3000]}"
            )
        })

# ------------------------------
# Upload de Arquivos
# ------------------------------
uploaded_files = st.file_uploader(
    "ðŸ“Ž Envie um ou mais arquivos (PDF ou DOCX)", type=["pdf", "docx"], accept_multiple_files=True
)

col1, col2 = st.columns([1, 1])

# ------------------------------
# Analisar contratos (Civil e BancÃ¡rio)
# ------------------------------
if area in ["Civil", "BancÃ¡rio"] and uploaded_files:
    with col1:
        if st.button("ðŸ” Analisar Contratos"):
            for up in uploaded_files:
                suffix = os.path.splitext(up.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(up.read())
                    caminho_temp = tmp.name

                try:
                    with st.spinner(f"Analisando {up.name}..."):
                        texto = ler_arquivo(caminho_temp)
                        analise = gerar_analise(texto, area=area)
                        st.session_state.analises[up.name] = analise
                        st.session_state.chat_history.append(
                            {"role": "assistant", "content": f"AnÃ¡lise do contrato {up.name}:\n{analise}"}
                        )
                except Exception as e:
                    st.error(f"Erro ao analisar {up.name}: {e}")

# ------------------------------
# Limpar anÃ¡lises
# ------------------------------
with col2:
    if st.session_state.analises and st.button("ðŸ—‘ Limpar anÃ¡lises"):
        st.session_state.analises = {}

# ------------------------------
# ExibiÃ§Ã£o das anÃ¡lises
# ------------------------------
if st.session_state.analises:
    st.subheader("ðŸ“‘ AnÃ¡lises geradas")

    for nome, analise in st.session_state.analises.items():
        with st.expander(f"{nome}"):
            st.text_area("Parecer JurÃ­dico", analise, height=350)

            st.download_button(
                "â¬‡ Baixar parecer",
                analise.encode("utf-8"),
                file_name=f"parecer_{nome}.txt",
                mime="text/plain"
            )

            if st.button(f"âœ Gerar PetiÃ§Ã£o Inicial para {nome}"):
                with st.spinner("Gerando petiÃ§Ã£o inicial..."):
                    try:
                        doc = Document()
                        doc.add_heading("PetiÃ§Ã£o Inicial", 0)
                        doc.add_paragraph(analise)
                        peticao_path = os.path.join(tempfile.gettempdir(), f"peticao_inicial_{nome}.docx")
                        doc.save(peticao_path)

                        with open(peticao_path, "rb") as f:
                            st.download_button(
                                f"â¬‡ Baixar PetiÃ§Ã£o ({nome})",
                                f,
                                file_name=f"peticao_inicial_{nome}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Erro ao gerar petiÃ§Ã£o: {e}")
else:
    st.info("Nenhuma anÃ¡lise disponÃ­vel ainda.")

# ------------------------------
# Chat JurÃ­dico
# ------------------------------
st.divider()
st.subheader(f"ðŸ’¬ Chat JurÃ­dico â€“ Ãrea: {area}")

# Exibe histÃ³rico
for msg in st.session_state.chat_history:
    if msg["role"] == "user":
        with st.chat_message("user"):
            st.markdown(msg["content"])
    elif msg["role"] == "assistant":
        with st.chat_message("assistant"):
            st.markdown(msg["content"])

# Entrada do chat
if prompt := st.chat_input("Digite sua pergunta para a IA jurÃ­dica"):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("Analisando..."):
            try:
                resposta = gerar_resposta_chat(st.session_state.chat_history)
                st.markdown(resposta)
                st.session_state.chat_history.append({"role": "assistant", "content": resposta})
            except Exception as e:
                st.error(f"Erro ao responder: {e}")

# ------------------------------
# Sidebar
# ------------------------------
with st.sidebar:
    st.header("âš™ï¸ ConfiguraÃ§Ãµes")
    if st.button("ðŸ§¹ Limpar histÃ³rico do chat"):
        st.session_state.chat_history = [{"role": "system", "content": get_prompt_sistema(area)}]
    st.markdown("---")
    st.caption("Desenvolvido com â¤ï¸ por vocÃª e IA")
>>>>>>> 84bc48d03e12871b7a7bf01d0c98fbf1f1bbbe91
